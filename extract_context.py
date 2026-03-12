import json
import os
from docx import Document

def main():
    json_path = os.path.join("dataset", "d1.json")
    docx_path = "contexts.docx"

    print(f"Loading data from {json_path}...")
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Create a new Word document
    doc = Document()
    count = 0
    
    # Iterate through the structure to extract context
    if 'data' in data:
        for article in data['data']:
            if 'title' in article:
                doc.add_heading(article['title'], level=1)
            if 'paragraphs' in article:
                for idx, paragraph in enumerate(article['paragraphs']):
                    if 'context' in paragraph:
                        doc.add_paragraph(paragraph['context'])
                        count += 1
                        
    print(f"Extracted {count} contexts.")
    
    # Save the document
    print(f"Saving to {docx_path}...")
    doc.save(docx_path)
    print("Done!")

if __name__ == "__main__":
    main()
